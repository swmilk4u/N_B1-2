# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(29.7)
section.page_height   = Cm(21.0)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(1.8)

NAVY       = (0x0D, 0x1B, 0x2A)
BLUE_DARK  = (0x1A, 0x53, 0x76)
GRAY_DARK  = (0x22, 0x22, 0x22)
GRAY_MID   = (0x77, 0x77, 0x77)
SANGWOO_C  = (0xD4, 0x8C, 0x00)   # 상우 : 황금빛 amber
YUJIN_C    = (0x00, 0x87, 0xBD)   # 유진 : 차가운 cyan-blue


def set_font(run, size=11, bold=False, color=None, name="malgun gothic"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_border(p, color_hex="1A5376"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "6")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color_hex)
    pBdr.append(b)
    pPr.append(pBdr)

def add_box_border(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "12")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "AAAAAA")
        pBdr.append(el)
    pPr.append(pBdr)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    sz  = {1:20, 2:15, 3:13}.get(level, 13)
    col = {1:NAVY, 2:BLUE_DARK, 3:BLUE_DARK}.get(level)
    run = p.add_run(text)
    set_font(run, size=sz, bold=True, color=col)
    add_border(p)

def body(doc, text, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, color=color or GRAY_DARK)

def shade(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cfont(cell, text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color or GRAY_DARK)

def table(doc, headers, rows, widths=None, hdr_color="1A5376"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, hdr_color)
        cfont(c, h, size=10, bold=True, color=(0xFF,0xFF,0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        bg = "F0F4F8" if ri%2==0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            shade(c, bg)
            cfont(c, val, size=10)
    if widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(widths):
                t.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

def prompt_box(doc, label, en, ko):
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    for ci in range(2):
        shade(t.cell(0,ci), "0D3B66")
    cfont(t.cell(0,0), f"PROMPT  {label}  (Copy & Paste)", size=10, bold=True, color=(0xFF,0xDA,0x79))
    cfont(t.cell(0,1), "Korean Summary", size=10, bold=True, color=(0xA8,0xD8,0xEA))
    shade(t.cell(1,0), "F8F9FA")
    shade(t.cell(1,1), "EFF6FB")
    cfont(t.cell(1,0), en, size=9.5)
    cfont(t.cell(1,1), ko, size=9.5, color=(0x33,0x33,0x33))
    for r in t.rows:
        r.cells[0].width = Cm(13.0)
        r.cells[1].width = Cm(11.5)
    doc.add_paragraph()

def char_tag(doc, who="BOTH"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if who == "ARA":
        txt, col = "[Sangwoo / 상우]  Amber warmth - male artist character", SANGWOO_C
    elif who == "YUJIN":
        txt, col = "[Yujin / 유진]  Cyan-blue - female AI creator character", YUJIN_C
    else:
        txt, col = "[Sangwoo + Yujin / 상우 + 유진]  Both characters on scene", BLUE_DARK
    run = p.add_run(txt)
    set_font(run, size=11, bold=True, color=col)

def style_note(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run("[Visual Style]  ")
    set_font(r1, size=10, bold=True, color=BLUE_DARK)
    r2 = p.add_run(
        "Background & environment : full desaturated B&W monochrome  /  "
        "Characters (Sangwoo/상우, Yujin/유진) and key objects : color only  /  "
        "Sangwoo (Male) = warm golden amber (#D48C00)  /  Yujin (Female) = cool cyan-blue (#0087BD)"
    )
    set_font(r2, size=10, color=GRAY_MID)

def tts_box(doc, text, note):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run("[ElevenLabs Narration]  ")
    set_font(r1, size=10, bold=True, color=BLUE_DARK)
    r2 = p.add_run(text)
    set_font(r2, size=12, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    r3 = p2.add_run(f"Voice style : {note}")
    set_font(r3, size=9.5, color=GRAY_MID)
    doc.add_paragraph()

def img_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"[IMAGE PLACEHOLDER]  {caption}")
    set_font(run, size=10, bold=True, color=(0x88,0x88,0x88))
    add_box_border(p)
    for _ in range(3):
        ep = doc.add_paragraph(" ")
        ep.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


# ===========================================================
# Cover Page
# ===========================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
set_font(p.add_run("Storyboard Planning Document"), size=28, bold=True, color=NAVY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p2.add_run("Campaign : Before You Create, Think Once More"), size=16, color=GRAY_MID)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p3.add_run("K-AI Contents Award 2026  |  Track A. University/General  AI Commercial"), size=12, color=GRAY_MID)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p4.add_run("2026-06-14"), size=11, color=(0x99,0x99,0x99))

doc.add_page_break()


# ===========================================================
# 1. Brand Identity
# ===========================================================
heading(doc, "1. Brand Identity  /  브랜드 아이덴티티")
table(doc,
    headers=["Item", "Content"],
    rows=[
        ["Campaign Name",    "Before You Create, Think Once More  /  만들기 전에, 한 번 더"],
        ["Contest Theme",    "AI Ethics - Etiquette & Copyright Protection in Generative AI Use"],
        ["Target Audience",  "Students, creators, and general public using generative AI tools (ages 20-40)"],
        ["Tone & Manner",    "Artistic & cinematic / B&W background + character color (Selective Color) / emotional depth"],
        ["Visual Concept",
         "Background & environment : fully desaturated black & white monochrome\n"
         "Sangwoo (Male/상우) + Yujin (Female/유진) + key objects : color only\n"
         "Sangwoo = warm golden amber #D48C00 (창작의 온기)  |  Yujin = cool cyan-blue #0087BD (digital/AI)"],
        ["USP",              "Treating 'thinking about creators before using AI' as a new creative culture"],
        ["Core Message",     "Even when creating with AI, we think of the original creator first."],
        ["Ending Slogan",    "Before you create, think once more.  /  만들기 전에, 한 번 더."],
    ],
    widths=[4.5, 20.0]
)


# ===========================================================
# 2. Character Profiles
# ===========================================================
heading(doc, "2. Character Profiles  /  등장 캐릭터 소개")
table(doc,
    headers=["Item", "Sangwoo (상우)  -  Original Creator", "Yujin  (유진)  -  AI Creator"],
    rows=[
        ["Role",
         "Traditional male artist / symbolizes original copyright holders",
         "Female digital creator using AI tools / symbolizes AI content creators"],
        ["Appearance",
         "Young Korean man in late 20s / warm-toned skin / dark brown wavy hair / worn artist apron, palette, fine brush",
         "Young Korean woman in mid-to-late 20s / clean casual style (turtleneck+wide trousers) / slim laptop, wireless earbuds"],
        ["Signature Color",
         "Warm golden amber  (#D48C00)\n창작의 열정과 온기를 상징",
         "Cool cyan-blue  (#0087BD)\n디지털·AI 세계를 상징하는 차가운 빛"],
        ["Scene Appearances",
         "Scene 1-A (hand closeup), Scene 2 (artwork particles), Scene 4 (coexistence), Scene 5 (ending)",
         "Scene 1-B (hand closeup), Scene 2 (work particles), Scene 3 (realization), Scene 4 (coexistence), Scene 5 (ending)"],
        ["Selective Color Rule",
         "Full figure, hands, apron, brush, palette rendered in amber / background B&W",
         "Full figure, laptop screen, AI icon rendered in cyan-blue / background B&W"],
    ],
    widths=[3.5, 11.5, 9.5],
    hdr_color="0D3B66"
)
body(doc,
    "Character Consistency Tips:\n"
    "- Include the same appearance description in each scene prompt to maintain visual consistency\n"
    "- If Gemini generates inconsistent looks: regenerate with the same prompt or use Runway Character Reference\n"
    "- In CapCut: apply selective color adjustment layer to strengthen the B&W background / color character effect"
)
doc.add_page_break()


# ===========================================================
# 3. Narrative Structure
# ===========================================================
heading(doc, "3. Narrative Structure  /  캠페인 목표 및 서사 구조")
body(doc, "Ad Objective : AI copyright & etiquette awareness  (recognition + attitude change)")
doc.add_paragraph()
table(doc,
    headers=["Structure", "Narrative", "Characters", "Scene"],
    rows=[
        ["Act 1 - Introduction",  "A work of art is born from Sangwoo's hands",                    "Sangwoo",         "Scene 1"],
        ["Act 2 - Rising Action", "Sangwoo and Yujin's works trained the AI",                        "Sangwoo + Yujin", "Scene 2"],
        ["Act 3 - Turning Point", "Yujin realizes she is also a creator",                            "Yujin",           "Scene 3"],
        ["Act 4 - Resolution",    "Sangwoo and Yujin side by side — cite, ask permission, respect",  "Sangwoo + Yujin", "Scene 4 + 5"],
    ],
    widths=[3.5, 10.5, 3.5, 2.0]
)


# ===========================================================
# 4. Video Spec
# ===========================================================
heading(doc, "4. Video Specification  /  영상 스펙")
table(doc,
    headers=["Item", "Spec"],
    rows=[
        ["Duration",       "40 seconds"],
        ["Resolution",     "1920 x 1080 (1080p)"],
        ["Aspect Ratio",   "16:9 Landscape"],
        ["Frame Rate",     "24fps"],
        ["Video Codec",    "H.264"],
        ["Audio Codec",    "AAC"],
        ["BGM",            "Contest official designated music only"],
        ["Visual Style",   "B&W background + character/key object color  (Selective Color)"],
        ["AI Watermark",   "Bottom-right corner as per official guidelines"],
    ],
    widths=[5.5, 19.0]
)


# ===========================================================
# 5. Tool List
# ===========================================================
heading(doc, "5. Tool List  /  사용 도구 목록")
table(doc,
    headers=["Purpose", "Primary Tool", "Backup Tool", "Reason"],
    rows=[
        ["Image Generation",  "Google Gemini",             "Ideogram / Canva AI",  "High prompt accessibility; supports both photorealistic and illustration styles"],
        ["Video Conversion",  "Runway Gen-3",              "Pika",                 "Excellent image-to-motion quality; precise motion control per scene"],
        ["Voice TTS",         "ElevenLabs",                "CLOVA Voice",          "Natural, emotional Korean voice support"],
        ["BGM",               "Contest official music",    "—",                    "Contest regulation requirement"],
        ["Editing",           "CapCut",                    "DaVinci Resolve",      "Text animation, AI watermark insertion, selective color effect, BGM fade-out"],
    ],
    widths=[3.5, 4.5, 4.5, 12.0]
)
doc.add_page_break()


# ===========================================================
# Scene 1
# ===========================================================
heading(doc, "6. Scene 1  -  Creator's Hands  /  창작자의 손  (0~7sec, 7sec)")
char_tag(doc, "BOTH")
style_note(doc)
table(doc,
    headers=["Field", "Content"],
    rows=[
        ["Scene No.",      "1"],
        ["Duration",       "7 seconds  (0~7sec)"],
        ["Objective",      "Visually imprint the fact that a work of art is born from a human hand"],
        ["Composition",    "Close-up / Sangwoo's brush hand (amber) + Yujin's typing hand (cyan) on B&W background / No text"],
        ["Narration",      "None  (music and visuals only)"],
        ["Tools",          "Image : Google Gemini  |  Video : Runway Gen-3"],
        ["Tool Reason",    "Generate 2 hand close-up keyvisuals with Gemini, convert to slow-motion with Runway"],
    ],
    widths=[5.5, 19.0]
)

heading(doc, "Scene 1-A : Sangwoo's Hand  /  상우의 손  (Painter)", level=2)
char_tag(doc, "ARA")
prompt_box(doc,
    label="Gemini Image Prompt",
    en=(
        "Selective color photography: close-up of a young Korean man's hand (Sangwoo) holding a fine brush,\n"
        "making a delicate stroke on a canvas.\n"
        "The hand, brush, and paint are rendered in warm golden amber color (#D48C00),\n"
        "everything else - studio background, canvas edges - is desaturated black and white monochrome.\n"
        "Dramatic chiaroscuro lighting, shallow depth of field, cinematic, photorealistic, 4K."
    ),
    ko=(
        "상우(남성)의 손·붓·물감 : 황금빛 amber 컬러만 유지\n"
        "나머지 배경·스튜디오 : 흑백\n"
        "명암 대비 강한 조명, 얕은 심도, 시네마틱"
    )
)
img_placeholder(doc, "Scene 1-A image  |  filename : scene01A_sangwoo_hand.png")
prompt_box(doc,
    label="Runway Motion Prompt",
    en=(
        "Sangwoo's golden amber hand slowly raises the brush and makes one deliberate stroke,\n"
        "the amber color glows warmly against the black and white background,\n"
        "cinematic slow motion, 24fps."
    ),
    ko="상우의 amber 색 붓이 흑백 배경 위에서 천천히 한 획을 긋는 슬로우모션"
)

heading(doc, "Scene 1-B : Yujin's Hand  /  유진의 손  (Typist)", level=2)
char_tag(doc, "YUJIN")
prompt_box(doc,
    label="Gemini Image Prompt",
    en=(
        "Selective color photography: close-up of a young person's hand (Yujin) typing on a keyboard,\n"
        "fingers mid-keystroke, one key in sharp focus.\n"
        "The hand, fingertips, and the glowing key are rendered in cool cyan-blue color (#0087BD),\n"
        "everything else - desk, keyboard body, background - is desaturated black and white monochrome.\n"
        "Dark moody background, cinematic, photorealistic, 4K, shallow depth of field."
    ),
    ko=(
        "유진의 손·손가락·키 : cyan-blue 컬러만 유지\n"
        "나머지 책상·키보드 몸체·배경 : 흑백\n"
        "어두운 분위기, 얕은 심도, 시네마틱"
    )
)
img_placeholder(doc, "Scene 1-B image  |  filename : scene01B_yujin_hand.png")
prompt_box(doc,
    label="Runway Motion Prompt",
    en=(
        "Yujin's cyan-blue fingers type rhythmically on the keyboard,\n"
        "each keystroke sends a subtle cyan ripple across the black and white surface,\n"
        "cinematic slow motion."
    ),
    ko="cyan 색 손가락이 흑백 키보드 위를 리드미컬하게 타이핑, 키마다 잔물결 효과"
)
doc.add_page_break()


# ===========================================================
# Scene 2
# ===========================================================
heading(doc, "7. Scene 2  -  AI Learning Visualization  /  AI 학습 시각화  (7~17sec, 10sec)")
char_tag(doc, "BOTH")
style_note(doc)
table(doc,
    headers=["Field", "Content"],
    rows=[
        ["Scene No.",   "2"],
        ["Duration",    "10 seconds  (7~17sec)"],
        ["Objective",   "Abstractly visualize that Sangwoo and Yujin's works trained the AI"],
        ["Composition", "Wide shot / B&W space background / amber particles (Sangwoo) + cyan particles (Yujin) converge into AI icon / narration text"],
        ["Narration",   '"All these works of art... created the AI."'],
        ["Tools",       "Image : Gemini  |  Video : Runway  |  Voice : ElevenLabs"],
        ["Tool Reason", "Two-color particle convergence visualizes coexistence; Runway VFX motion; ElevenLabs emotional narration"],
    ],
    widths=[5.5, 19.0]
)
prompt_box(doc,
    label="Gemini Image Prompt",
    en=(
        "Selective color abstract visualization: black and white dark space background.\n"
        "From the left side, warm golden amber particles (representing Sangwoo's artworks -\n"
        "tiny paintings, brushstrokes, sketches) stream inward.\n"
        "From the right side, cool cyan-blue particles (representing Yujin's digital works -\n"
        "code snippets, pixel art, AI-generated thumbnails) stream inward.\n"
        "Both particle streams converge and merge into a single glowing white AI neural network icon\n"
        "at the center of the frame.\n"
        "High contrast, cinematic VFX, dramatic, 4K."
    ),
    ko=(
        "흑백 우주 배경\n"
        "왼쪽 : 상우 작품(그림·붓터치) -> amber 파티클\n"
        "오른쪽 : 유진 작품(코드·AI 이미지) -> cyan 파티클\n"
        "두 파티클이 중앙 빛나는 AI 아이콘으로 수렴·합쳐짐"
    )
)
img_placeholder(doc, "Scene 2 image  |  filename : scene02_ai_learning.png")
prompt_box(doc,
    label="Runway Motion Prompt",
    en=(
        "Amber and cyan particle streams swirl from both sides and merge into the central AI icon,\n"
        "the icon pulses with white light as both colors blend inside it,\n"
        "smooth cinematic VFX motion, dramatic convergence."
    ),
    ko="amber·cyan 파티클이 양쪽에서 수렴해 AI 아이콘 안에서 합쳐지며 흰빛으로 맥동"
)
tts_box(doc,
    text='"All these works of art...  created the AI."  /  "이 모든 창작물이...  AI를 만들었습니다."',
    note="Calm, low tone / slow pace / emotional"
)
doc.add_page_break()


# ===========================================================
# Scene 3
# ===========================================================
heading(doc, "8. Scene 3  -  Yujin's Realization  /  유진의 깨달음  (17~25sec, 8sec)")
char_tag(doc, "YUJIN")
style_note(doc)
table(doc,
    headers=["Field", "Content"],
    rows=[
        ["Scene No.",   "3"],
        ["Duration",    "8 seconds  (17~25sec)"],
        ["Objective",   "Help viewers realize, alongside Yujin, that AI users are also creators"],
        ["Composition", "Medium shot / B&W room / Yujin (cyan) at laptop / screen emits amber+cyan light / narration text"],
        ["Narration",   '"Now, you are also a creator."'],
        ["Tools",       "Image : Gemini  |  Video : Runway  |  Voice : ElevenLabs"],
        ["Tool Reason", "Yujin as color subject invites viewer self-projection; laptop screen light transitions from dark (Scene 2) to hopeful"],
    ],
    widths=[5.5, 19.0]
)
prompt_box(doc,
    label="Gemini Image Prompt",
    en=(
        "Selective color photography: medium shot of Yujin, a young person in casual clothes,\n"
        "sitting at a desk with a laptop in a room.\n"
        "Yujin's entire figure is rendered in cool cyan-blue tones (#0087BD),\n"
        "the laptop screen glows with a mix of amber and cyan light (showing AI-generated artwork).\n"
        "Everything else - room walls, desk, chair, background - is fully desaturated black and white.\n"
        "Yujin slowly raises their head and looks at the screen with a sense of realization.\n"
        "Cinematic, hopeful atmosphere, soft backlight from window (black and white), 4K."
    ),
    ko=(
        "유진 전신 : cyan-blue 컬러\n"
        "노트북 화면 : amber+cyan 빛\n"
        "방·책상·배경 : 흑백\n"
        "유진이 고개를 들어 화면을 바라보는 순간 (깨달음)"
    )
)
img_placeholder(doc, "Scene 3 image  |  filename : scene03_yujin_realization.png")
prompt_box(doc,
    label="Runway Motion Prompt",
    en=(
        "Yujin slowly raises their head and looks at the glowing laptop screen,\n"
        "the cyan glow around Yujin intensifies slightly as realization sets in,\n"
        "camera gently pushes in toward Yujin's face and screen,\n"
        "cinematic slow motion, hopeful tone."
    ),
    ko="유진이 고개를 들어 화면을 바라보며 cyan 빛이 강해짐, 카메라가 천천히 당겨짐"
)
tts_box(doc,
    text='"Now,  you are also a creator."  /  "이제,  당신도 창작자입니다."',
    note="Warm, gentle tone / confident pace"
)
doc.add_page_break()


# ===========================================================
# Scene 4
# ===========================================================
heading(doc, "9. Scene 4  -  Coexistence & Three Actions  /  공존과 행동 지침  (25~35sec, 10sec)")
char_tag(doc, "BOTH")
style_note(doc)
table(doc,
    headers=["Field", "Content"],
    rows=[
        ["Scene No.",   "4"],
        ["Duration",    "10 seconds  (25~35sec)"],
        ["Objective",   "Sangwoo and Yujin stand together and clearly deliver three actions: cite, ask permission, respect"],
        ["Composition", "Center frame / B&W background / Sangwoo (amber) + Yujin (cyan) side by side / white light blooms between them / 3-step text fade-in"],
        ["Narration",   '"Cite the source, ask for permission, respect the creator."'],
        ["Tools",       "Image : Gemini  |  Video : Runway  |  Voice : ElevenLabs  |  Edit : CapCut"],
        ["Tool Reason", "Two-color character contrast visualizes coexistence; CapCut handles 3-step text fade-in"],
    ],
    widths=[5.5, 19.0]
)
prompt_box(doc,
    label="Gemini Image Prompt",
    en=(
        "Selective color illustration: two characters standing side by side on a clean background.\n"
        "Left - Sangwoo: a young Korean man in an artist's apron holding a paintbrush and palette,\n"
        "rendered entirely in warm golden amber color (#D48C00).\n"
        "Right - Yujin: a young Korean woman in casual modern clothes holding a laptop,\n"
        "rendered entirely in cool cyan-blue color (#0087BD).\n"
        "Between them, a soft white glowing light emanates upward.\n"
        "The entire background is desaturated black and white monochrome.\n"
        "Minimalist style, balanced composition, no text, cinematic lighting, 4K."
    ),
    ko=(
        "흑백 배경\n"
        "왼쪽 상우(amber) : 앞치마·붓·팔레트\n"
        "오른쪽 유진(cyan) : 노트북\n"
        "두 사람 사이에서 흰빛이 위로 퍼짐\n"
        "텍스트 없음, 균형 잡힌 구도"
    )
)
img_placeholder(doc, "Scene 4 image  |  filename : scene04_coexist.png")
prompt_box(doc,
    label="Runway Motion Prompt",
    en=(
        "Sangwoo and Yujin simultaneously turn to face each other and nod gently,\n"
        "the white light between them blooms and pulses warmly,\n"
        "amber and cyan glows intensify slightly, smooth cinematic animation."
    ),
    ko="상우·유진이 서로를 향해 고개를 끄덕이고, 사이의 흰빛이 따뜻하게 번짐"
)
tts_box(doc,
    text='"Cite the source,  ask for permission,  respect the creator."  /  "출처를 밝히고,  허락을 구하고,  창작자를 존중합니다."',
    note="Clear, articulate tone / slight pause between three phrases"
)
heading(doc, "CapCut Text Editing  (Scene 4)", level=3)
table(doc,
    headers=["Timing", "Text (Korean)", "Text (English)", "Effect"],
    rows=[
        ["0~3sec",  "출처를 밝히고,",       "Cite the source,",        "Fade in"],
        ["3~6sec",  "허락을 구하고,",        "Ask for permission,",     "Fade in"],
        ["6~10sec", "창작자를 존중합니다.",  "Respect the creator.",    "Fade in + Hold"],
    ],
    widths=[2.0, 5.5, 7.0, 3.5]
)
doc.add_page_break()


# ===========================================================
# Scene 5
# ===========================================================
heading(doc, "10. Scene 5  -  Ending Slogan Card  /  엔딩 카드  (35~40sec, 5sec)")
char_tag(doc, "BOTH")
table(doc,
    headers=["Field", "Content"],
    rows=[
        ["Scene No.",   "5"],
        ["Duration",    "5 seconds  (35~40sec)"],
        ["Objective",   "Imprint the slogan and close with AI watermark and brand identity"],
        ["Composition", "B&W background / Sangwoo (amber) + Yujin (cyan) silhouettes on left & right / slogan white text center"],
        ["Narration",   '"Before you create, think once more."  /  "만들기 전에, 한 번 더."'],
        ["Tools",       "Image : Gemini (optional)  |  Edit : CapCut  |  Voice : ElevenLabs"],
        ["Image Prompt","Optional - see below"],
        ["Output File", "scene05_ending.mp4"],
    ],
    widths=[5.5, 19.0]
)
prompt_box(doc,
    label="Gemini Image Prompt  (Optional)",
    en=(
        "Selective color: minimalist ending card.\n"
        "Black and white background.\n"
        "Far left: Sangwoo's silhouette in warm amber color.\n"
        "Far right: Yujin's silhouette in cool cyan-blue color.\n"
        "Center: empty space for text overlay.\n"
        "A thin horizontal line of white light connects the two silhouettes.\n"
        "Cinematic, minimal, 4K."
    ),
    ko=(
        "흑백 배경\n"
        "왼쪽 상우 실루엣(amber) / 오른쪽 유진 실루엣(cyan)\n"
        "중앙 : 텍스트 오버레이 공간\n"
        "두 실루엣을 잇는 흰 수평선"
    )
)
img_placeholder(doc, "Scene 5 ending background  |  filename : scene05_ending_bg.png")
heading(doc, "CapCut Settings", level=3)
table(doc,
    headers=["Item", "Value"],
    rows=[
        ["Background",    "B&W with Sangwoo(amber) + Yujin(cyan) silhouette image"],
        ["Main Text",     "Before you create, think once more.  /  만들기 전에, 한 번 더."],
        ["Text Color",    "#FFFFFF (white)"],
        ["Font",          "Noto Sans KR Bold or equivalent"],
        ["Sub Text",      "Even with AI, we think of the creator first.  /  AI로 만들 때도, 창작자를 먼저 생각합니다."],
        ["Appear Effect", "Fade in  (0.5 sec)"],
        ["BGM",           "Contest official music -> fade out"],
        ["AI Watermark",  "Bottom-right corner as per official guidelines"],
    ],
    widths=[5.5, 19.0]
)
tts_box(doc,
    text='"Before you create,  think once more."  /  "만들기 전에,  한 번 더."',
    note="Low, quiet tone with lingering feeling / slow pace"
)
doc.add_page_break()


# ===========================================================
# Prompt Revision Log
# ===========================================================
heading(doc, "11. Prompt Revision Log  /  프롬프트 수정 전/후 기록  (Scene 2)")
table(doc,
    headers=["Item", "Content"],
    rows=[
        ["Before (v1) Prompt",
         "many images and text flowing into an AI icon"],
        ["Problem with v1",
         "Generated in simple cartoon style. No B&W/color separation. "
         "Lacked scale and tension of particles. Could not distinguish Sangwoo vs Yujin works."],
        ["Revision Direction",
         "1. Specify Selective Color technique (amber/cyan separation)\n"
         "2. Specify B&W background explicitly (black and white dark space background)\n"
         "3. Assign direction and color to each particle stream (left=amber=Sangwoo / right=cyan=Yujin)\n"
         "4. Specify convergence point and merge behavior (merge into glowing white AI icon)"],
        ["After (v2) Prompt",
         "Selective color abstract visualization: black and white dark space background. "
         "From the left, warm golden amber particles (Sangwoo's artworks) stream inward. "
         "From the right, cool cyan-blue particles (Yujin's digital works) stream inward. "
         "Both converge and merge into a single glowing white AI neural network icon at center. "
         "High contrast, cinematic VFX, dramatic, 4K."],
        ["Result",
         "Obtained amber/cyan dual-particle convergence image on B&W background. "
         "Color separation clearly distinguishes Sangwoo and Yujin's works. Ad tone & manner matched."],
    ],
    widths=[5.5, 19.0]
)


# ===========================================================
# Timeline Summary
# ===========================================================
heading(doc, "12. Timeline Summary  /  타임라인 요약")
table(doc,
    headers=["Scene", "Time", "Dur.", "Key Content", "Characters", "Tools"],
    rows=[
        ["1", "0~7s",   "7s",  "Creator's hands : Sangwoo (amber) + Yujin (cyan) 2-cut",    "Sangwoo + Yujin", "Gemini + Runway"],
        ["2", "7~17s",  "10s", "Amber/cyan particles converge into AI icon + narration",    "Sangwoo + Yujin", "Gemini + Runway + ElevenLabs"],
        ["3", "17~25s", "8s",  "Yujin's realization (cyan) + narration",                    "Yujin",           "Gemini + Runway + ElevenLabs"],
        ["4", "25~35s", "10s", "Coexistence + cite, permission, respect 3-step",            "Sangwoo + Yujin", "Gemini + Runway + ElevenLabs + CapCut"],
        ["5", "35~40s", "5s",  "Ending slogan card + AI watermark",                         "Sangwoo + Yujin", "ElevenLabs + CapCut"],
        ["Total", "", "40s",   "", "", ""],
    ],
    widths=[1.2, 2.2, 1.5, 9.0, 3.5, 7.1]
)


# ===========================================================
# Save
# ===========================================================
OUTPUT = (
    r"c:\Users\YSW1\Desktop\20260601 codyssey_mission"
    r"\N_B1-2_Contents\01_document\storyboard_v2.docx"
)
doc.save(OUTPUT)
print(f"Saved : {OUTPUT}")
