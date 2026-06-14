# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21.0)   # A4 세로
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── 색상 ──────────────────────────────────────────────────
NAVY        = (0x0D, 0x1B, 0x2A)
BLUE_DARK   = (0x1A, 0x53, 0x76)
GRAY_DARK   = (0x22, 0x22, 0x22)
GRAY_MID    = (0x77, 0x77, 0x77)
WHITE       = (0xFF, 0xFF, 0xFF)
SANGWOO_C   = (0xD4, 0x8C, 0x00)   # 상우 : 황금빛 amber
YUJIN_C     = (0x00, 0x87, 0xBD)   # 유진 : 차가운 cyan-blue


# ── 헬퍼 ──────────────────────────────────────────────────
def sf(run, size=11, bold=False, color=None, name="malgun gothic"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cfont(cell, text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    sf(run, size=size, bold=bold, color=color or GRAY_DARK)

def set_col_width(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)

def img_box(doc, caption, height_lines=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"[ IMAGE : {caption} ]")
    sf(run, size=10, bold=True, color=(0x88, 0x88, 0x88))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top","left","bottom","right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "12")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "BBBBBB")
        pBdr.append(el)
    pPr.append(pBdr)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    for _ in range(height_lines):
        ep = doc.add_paragraph(" ")
        ep.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()

def prompt_box(doc, label, en, ko, accent_hex):
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    for ci in range(2):
        shade(t.cell(0, ci), accent_hex)
    cfont(t.cell(0,0), f"PROMPT  {label}", size=9, bold=True, color=(0xFF,0xFF,0xFF))
    cfont(t.cell(0,1), "Korean Summary", size=9, bold=True, color=(0xFF,0xFF,0xFF))
    shade(t.cell(1,0), "F8F9FA")
    shade(t.cell(1,1), "EFF6FB")
    cfont(t.cell(1,0), en,  size=9)
    cfont(t.cell(1,1), ko,  size=9, color=(0x33,0x33,0x33))
    set_col_width(t, [8.8, 6.0])
    doc.add_paragraph()

def divider(doc, color_hex):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "8")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color_hex)
    pBdr.append(b)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def section_title(doc, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sf(run, size=14, bold=True, color=color)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "6")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), "%02X%02X%02X" % color)
    pBdr.append(b)
    pPr.append(pBdr)

def info_table(doc, rows, widths, header_hex):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for ri, (lbl, val) in enumerate(rows):
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        shade(t.rows[ri].cells[0], header_hex)
        shade(t.rows[ri].cells[1], bg)
        cfont(t.rows[ri].cells[0], lbl, size=10, bold=True, color=WHITE)
        cfont(t.rows[ri].cells[1], val, size=10)
    set_col_width(t, widths)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
sf(p.add_run("Character Reference Sheet"), size=26, bold=True, color=NAVY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p2.add_run("Campaign  :  Before You Create, Think Once More"), size=14, color=GRAY_MID)
sf(p2.add_run("  /  만들기 전에, 한 번 더"), size=14, color=GRAY_MID)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p3.add_run("K-AI Contents Award 2026  |  Track A. AI Commercial"), size=11, color=GRAY_MID)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p4.add_run("2026-06-14"), size=10, color=(0x99,0x99,0x99))

doc.add_paragraph()

# 캐릭터 2명 컬러 요약
t_sum = doc.add_table(rows=1, cols=2)
t_sum.style = "Table Grid"
shade(t_sum.cell(0,0), "D48C00")
shade(t_sum.cell(0,1), "0087BD")
cfont(t_sum.cell(0,0),
      "SANGWOO  (상우)\nOriginal Creator / Artist\nSignature Color : Warm Golden Amber  #D48C00",
      size=12, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
cfont(t_sum.cell(0,1),
      "YUJIN  (유진)\nAI Creator / Digital Artist\nSignature Color : Cool Cyan-Blue  #0087BD",
      size=12, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
set_col_width(t_sum, [8.5, 8.5])
doc.add_paragraph()

doc.add_page_break()


# ══════════════════════════════════════════════════════════
#  CHARACTER 1 : SANGWOO (상우)
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
sf(p.add_run("CHARACTER  01"), size=11, color=GRAY_MID)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
sf(p2.add_run("SANGWOO  (상우)"), size=28, bold=True, color=SANGWOO_C)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p3.add_run("Original Creator  /  원작 창작자  |  Traditional Artist"), size=13, color=GRAY_MID)

divider(doc, "D48C00")
doc.add_paragraph()

# ── 상단 : 이미지 자리 + 기본 정보 나란히 ──────────────
t_top = doc.add_table(rows=1, cols=2)
t_top.style = "Table Grid"
shade(t_top.cell(0,0), "FFFBF0")
shade(t_top.cell(0,1), "FFFBF0")

# 왼쪽 : 전신 이미지
left_cell = t_top.cell(0,0)
left_cell.text = ""
lp = left_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(lp.add_run(
    "\n\n\n\n\n\n\n\n\n\n\n\n"
    "[ Full-body reference image ]\n"
    "scene_sangwoo_fullbody.png\n\n\n\n\n\n\n\n\n\n\n\n"
), size=9, color=(0xAA,0xAA,0xAA))

# 오른쪽 : 기본 정보
right_cell = t_top.cell(0,1)
right_cell.text = ""
for label, value in [
    ("Name",             "Sangwoo  /  상우"),
    ("Age",              "Late 20s"),
    ("Gender",           "Male  /  남성"),
    ("Role",             "Original creator / Traditional artist\n원작 창작자 / 전통 방식 아티스트"),
    ("Personality",      "Warm, dedicated, principled\n따뜻하고 성실하며 원칙을 중요시함"),
    ("Signature Color",  "Warm Golden Amber  #D48C00\n창작의 열정과 온기를 상징"),
    ("Selective Color",  "His figure, hands, tools rendered in amber\nBackground always B&W monochrome"),
    ("Scene",            "Scene 1-A / Scene 2 / Scene 4 / Scene 5"),
]:
    p = right_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}  ")
    sf(r1, size=10, bold=True, color=SANGWOO_C)
    r2 = p.add_run(value)
    sf(r2, size=10, color=GRAY_DARK)

set_col_width(t_top, [7.5, 9.5])
doc.add_paragraph()


# ── 외형 상세 설명 ─────────────────────────────────────
section_title(doc, "Appearance Details  /  외형 상세", SANGWOO_C)
info_table(doc,
    rows=[
        ("Body Type",    "Lean and moderately built / average height (175cm) / upright posture"),
        ("Face",         "Warm-toned skin / defined jawline / calm and sincere eyes / slightly messy hair"),
        ("Hair",         "Dark brown, slightly wavy / medium length / loosely pushed back"),
        ("Clothing",     "Worn artist apron over a plain white or cream linen shirt\nRolled-up sleeves, paint stains on apron"),
        ("Hands",        "Slender, expressive fingers / often holding a fine brush or pencil"),
        ("Key Props",    "Fine brush / wooden palette with paint / sketchbook"),
        ("Amber Apply",  "Entire figure including clothing, apron, skin, hair rendered in warm amber tones\n"
                         "Brush and palette glow amber / paint strokes glow amber on canvas"),
    ],
    widths=[4.0, 13.0],
    header_hex="%02X%02X%02X" % SANGWOO_C
)


# ── 표정 & 감정 가이드 ────────────────────────────────
section_title(doc, "Expression Guide  /  표정 & 감정 가이드", SANGWOO_C)
info_table(doc,
    rows=[
        ("Scene 1-A",  "Focused and serene / eyes on canvas / gentle grip on brush"),
        ("Scene 2",    "Unaware (shown only as particle source) — no direct expression"),
        ("Scene 4",    "Soft smile / turns toward Yujin / nods gently / warm eye contact"),
        ("Scene 5",    "Calm, dignified forward gaze / slight upward lift at corners of mouth"),
    ],
    widths=[3.0, 14.0],
    header_hex="%02X%02X%02X" % SANGWOO_C
)


# ── 캐릭터 일관성 유지 프롬프트 ──────────────────────
section_title(doc, "Character Consistency Prompt  /  캐릭터 고정 프롬프트", SANGWOO_C)
prompt_box(doc,
    label="Sangwoo  -  Full Body Reference",
    en=(
        "Selective color character illustration: a young Korean man in his late 20s named Sangwoo.\n"
        "Lean build, 175cm, upright posture. Dark brown slightly wavy medium-length hair pushed back.\n"
        "Warm-toned skin, calm and sincere eyes, defined jawline.\n"
        "Wearing a worn artist apron over a cream linen shirt with rolled-up sleeves and paint stains.\n"
        "Holding a fine brush in his right hand, wooden palette in his left.\n"
        "Entire figure rendered in warm golden amber color (#D48C00) —\n"
        "hair, skin, clothing, apron, brush, palette all in amber tones.\n"
        "Background fully desaturated black and white monochrome.\n"
        "Full-body shot, front-facing, cinematic lighting, clean minimal background, 4K."
    ),
    ko=(
        "상우 : 20대 후반 한국 남성\n"
        "마른 체형, 175cm, 바른 자세\n"
        "어두운 갈색 약간 웨이브 머리 (중간 길이, 뒤로 넘김)\n"
        "따뜻한 피부톤, 차분하고 성실한 눈, 뚜렷한 턱선\n"
        "낡은 앞치마 + 크림 린넨 셔츠, 소매 걷어올림, 물감 자국\n"
        "오른손 : 붓 / 왼손 : 나무 팔레트\n"
        "전신 amber 컬러 / 배경 흑백"
    ),
    accent_hex="%02X%02X%02X" % SANGWOO_C
)

img_box(doc, "Sangwoo Full-body Reference  |  filename : char_sangwoo_fullbody.png", height_lines=12)

# ── 씬별 등장 프롬프트 ───────────────────────────────
section_title(doc, "Scene Prompts  /  씬별 등장 프롬프트", SANGWOO_C)

prompt_box(doc,
    label="Scene 1-A  :  Hand Close-up",
    en=(
        "Selective color photography: extreme close-up of Sangwoo's hand holding a fine brush,\n"
        "making a single deliberate stroke on a white canvas.\n"
        "His hand, brush, and amber-colored paint stroke are rendered in warm golden amber (#D48C00).\n"
        "Everything else - studio background, canvas surface - is desaturated black and white.\n"
        "Chiaroscuro lighting, shallow depth of field, cinematic, photorealistic, 4K."
    ),
    ko=(
        "상우 손·붓·붓터치 : amber 컬러\n"
        "배경·캔버스 : 흑백\n"
        "명암 대비 강한 조명, 얕은 심도, 시네마틱"
    ),
    accent_hex="%02X%02X%02X" % SANGWOO_C
)
img_box(doc, "Scene 1-A  |  filename : scene01A_sangwoo_hand.png", height_lines=8)

prompt_box(doc,
    label="Scene 4  :  Coexistence  (with Yujin)",
    en=(
        "Selective color illustration: Sangwoo (left) and Yujin (right) standing side by side.\n"
        "Sangwoo: young Korean man in artist apron, holding brush and palette,\n"
        "entire figure rendered in warm golden amber (#D48C00).\n"
        "Yujin: young Korean woman in casual modern clothing, holding a laptop,\n"
        "entire figure rendered in cool cyan-blue (#0087BD).\n"
        "A soft white glowing light emanates upward between them.\n"
        "Background is fully desaturated black and white monochrome.\n"
        "Minimalist, balanced composition, no text, cinematic lighting, 4K."
    ),
    ko=(
        "상우(왼쪽, amber) + 유진(오른쪽, cyan) 나란히\n"
        "두 사람 사이에서 흰빛이 위로 퍼짐\n"
        "배경 흑백, 텍스트 없음, 균형 잡힌 구도"
    ),
    accent_hex="%02X%02X%02X" % SANGWOO_C
)
img_box(doc, "Scene 4  |  filename : scene04_coexist.png", height_lines=8)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
#  CHARACTER 2 : YUJIN (유진)
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
sf(p.add_run("CHARACTER  02"), size=11, color=GRAY_MID)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
sf(p2.add_run("YUJIN  (유진)"), size=28, bold=True, color=YUJIN_C)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p3.add_run("AI Creator  /  AI 크리에이터  |  Digital Artist"), size=13, color=GRAY_MID)

divider(doc, "0087BD")
doc.add_paragraph()

# ── 상단 : 이미지 자리 + 기본 정보 나란히 ──────────────
t_top2 = doc.add_table(rows=1, cols=2)
t_top2.style = "Table Grid"
shade(t_top2.cell(0,0), "F0F8FF")
shade(t_top2.cell(0,1), "F0F8FF")

left_cell2 = t_top2.cell(0,0)
left_cell2.text = ""
lp2 = left_cell2.paragraphs[0]
lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(lp2.add_run(
    "\n\n\n\n\n\n\n\n\n\n\n\n"
    "[ Full-body reference image ]\n"
    "scene_yujin_fullbody.png\n\n\n\n\n\n\n\n\n\n\n\n"
), size=9, color=(0xAA,0xAA,0xAA))

right_cell2 = t_top2.cell(0,1)
right_cell2.text = ""
for label, value in [
    ("Name",             "Yujin  /  유진"),
    ("Age",              "Mid-to-late 20s"),
    ("Gender",           "Female  /  여성"),
    ("Role",             "AI creator / Digital artist\nAI 도구를 활용하는 디지털 크리에이터"),
    ("Personality",      "Curious, quick-witted, open-minded\n호기심 많고 빠른 판단력, 열린 사고방식"),
    ("Signature Color",  "Cool Cyan-Blue  #0087BD\n디지털·AI 세계를 상징하는 차가운 빛"),
    ("Selective Color",  "Her figure, laptop, AI glow rendered in cyan-blue\nBackground always B&W monochrome"),
    ("Scene",            "Scene 1-B / Scene 2 / Scene 3 / Scene 4 / Scene 5"),
]:
    p = right_cell2.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}  ")
    sf(r1, size=10, bold=True, color=YUJIN_C)
    r2 = p.add_run(value)
    sf(r2, size=10, color=GRAY_DARK)

set_col_width(t_top2, [7.5, 9.5])
doc.add_paragraph()


# ── 외형 상세 설명 ─────────────────────────────────────
section_title(doc, "Appearance Details  /  외형 상세", YUJIN_C)
info_table(doc,
    rows=[
        ("Body Type",    "Slender, average height (163cm) / light and quick movements / confident posture"),
        ("Face",         "Cool-toned bright skin / sharp, intelligent eyes / soft but defined features"),
        ("Hair",         "Dark black hair / straight, shoulder-length / half-up or down with clean cut"),
        ("Clothing",     "Clean modern casual : fitted white or light grey turtleneck + wide-leg trousers\n"
                         "Minimal accessories / wireless earbuds often worn"),
        ("Hands",        "Neat, nimble fingers / often typing or touching laptop screen"),
        ("Key Props",    "Slim laptop (lid showing subtle AI interface glow) / wireless earbuds / phone"),
        ("Cyan Apply",   "Entire figure including clothing, skin, hair, and laptop rendered in cool cyan-blue tones\n"
                         "Laptop screen glows cyan-blue / AI icons on screen glow cyan"),
    ],
    widths=[4.0, 13.0],
    header_hex="%02X%02X%02X" % YUJIN_C
)


# ── 표정 & 감정 가이드 ────────────────────────────────
section_title(doc, "Expression Guide  /  표정 & 감정 가이드", YUJIN_C)
info_table(doc,
    rows=[
        ("Scene 1-B",  "Focused / eyes watching the keyboard / rhythmic typing movement"),
        ("Scene 2",    "Unaware (shown only as particle source) — no direct expression"),
        ("Scene 3",    "Starting with a downward gaze / slowly raises head / eyes widen slightly with realization\n"
                       "Lips part softly / a quiet, hopeful expression settles on the face"),
        ("Scene 4",    "Warm smile / turns toward Sangwoo / nods gently / bright confident eyes"),
        ("Scene 5",    "Calm, forward gaze / slight smile / dignified and hopeful"),
    ],
    widths=[3.0, 14.0],
    header_hex="%02X%02X%02X" % YUJIN_C
)


# ── 캐릭터 일관성 유지 프롬프트 ──────────────────────
section_title(doc, "Character Consistency Prompt  /  캐릭터 고정 프롬프트", YUJIN_C)
prompt_box(doc,
    label="Yujin  -  Full Body Reference",
    en=(
        "Selective color character illustration: a young Korean woman in her mid-to-late 20s named Yujin.\n"
        "Slender build, 163cm, confident and light posture. Dark black straight shoulder-length hair.\n"
        "Cool-toned bright skin, sharp intelligent eyes, soft but defined features.\n"
        "Wearing a fitted white turtleneck and wide-leg trousers.\n"
        "Holding a slim laptop in both hands, screen glowing softly.\n"
        "Wireless earbuds in ears.\n"
        "Entire figure rendered in cool cyan-blue color (#0087BD) —\n"
        "hair, skin, clothing, laptop, earbuds all in cyan-blue tones.\n"
        "Laptop screen glows with cyan-blue AI interface light.\n"
        "Background fully desaturated black and white monochrome.\n"
        "Full-body shot, front-facing, cinematic lighting, clean minimal background, 4K."
    ),
    ko=(
        "유진 : 20대 중후반 한국 여성\n"
        "슬렌더 체형, 163cm, 자신감 있는 가벼운 자세\n"
        "검은 직모 단발 (어깨 길이)\n"
        "차가운 밝은 피부톤, 날카롭고 지적인 눈, 부드럽고 뚜렷한 이목구비\n"
        "흰 터틀넥 + 와이드 트라우저 / 무선 이어폰\n"
        "양손에 슬림 노트북, 화면에서 cyan 빛\n"
        "전신 cyan-blue 컬러 / 배경 흑백"
    ),
    accent_hex="%02X%02X%02X" % YUJIN_C
)

img_box(doc, "Yujin Full-body Reference  |  filename : char_yujin_fullbody.png", height_lines=12)


# ── 씬별 등장 프롬프트 ───────────────────────────────
section_title(doc, "Scene Prompts  /  씬별 등장 프롬프트", YUJIN_C)

prompt_box(doc,
    label="Scene 1-B  :  Hand Close-up",
    en=(
        "Selective color photography: extreme close-up of Yujin's hand typing on a slim keyboard,\n"
        "fingers mid-keystroke, one key in sharp focus.\n"
        "Her hand, fingertips, and the glowing key are rendered in cool cyan-blue (#0087BD).\n"
        "Everything else - desk surface, keyboard body - is desaturated black and white.\n"
        "Dark moody background, cinematic, photorealistic, 4K, shallow depth of field."
    ),
    ko=(
        "유진 손·손가락·키 : cyan-blue 컬러\n"
        "책상·키보드 몸체·배경 : 흑백\n"
        "어두운 분위기, 얕은 심도, 시네마틱"
    ),
    accent_hex="%02X%02X%02X" % YUJIN_C
)
img_box(doc, "Scene 1-B  |  filename : scene01B_yujin_hand.png", height_lines=8)

prompt_box(doc,
    label="Scene 3  :  Realization Moment",
    en=(
        "Selective color photography: medium shot of Yujin sitting at a desk with a slim laptop.\n"
        "Yujin's entire figure is rendered in cool cyan-blue (#0087BD).\n"
        "The laptop screen glows with mixed amber and cyan light (AI-generated artwork visible).\n"
        "Everything else - room, desk, chair, walls - is fully desaturated black and white.\n"
        "Yujin starts with head slightly bowed, then slowly raises her head.\n"
        "Eyes widen gently with realization — lips part softly, a quiet hopeful expression.\n"
        "Soft backlight from window (black and white), cinematic, hopeful atmosphere, 4K."
    ),
    ko=(
        "유진 전신 : cyan-blue 컬러\n"
        "노트북 화면 : amber+cyan 혼합 빛\n"
        "방·책상·배경 : 흑백\n"
        "처음엔 고개 살짝 숙임 → 천천히 들어올림\n"
        "눈이 살짝 커지는 깨달음의 표정 / 희망적 분위기"
    ),
    accent_hex="%02X%02X%02X" % YUJIN_C
)
img_box(doc, "Scene 3  |  filename : scene03_yujin_realization.png", height_lines=8)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
#  VISUAL STYLE GUIDE
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
sf(p.add_run("VISUAL STYLE GUIDE"), size=22, bold=True, color=NAVY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p2.add_run("Selective Color  :  B&W Background + Character Color"), size=13, color=GRAY_MID)
divider(doc, "1A5376")
doc.add_paragraph()

t_style = doc.add_table(rows=1, cols=3)
t_style.style = "Table Grid"

shade(t_style.cell(0,0), "222222")
shade(t_style.cell(0,1), "D48C00")
shade(t_style.cell(0,2), "0087BD")

cfont(t_style.cell(0,0),
      "BACKGROUND\n\nFully desaturated\nBlack & White\nMonochrome\n\n"
      "배경·환경·소품\n모두 흑백 처리\n\n"
      "Effect achieved via\nCapCut selective\ncolor adjustment layer",
      size=10, bold=False, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

cfont(t_style.cell(0,1),
      "SANGWOO  (상우)\n\nWarm Golden Amber\n#D48C00\n\n"
      "상우 : 전신·손·도구\n모두 amber 컬러\n\n"
      "창작의 열정과\n온기를 상징\n\n"
      "Brushstroke glow :\namber trail on B&W canvas",
      size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

cfont(t_style.cell(0,2),
      "YUJIN  (유진)\n\nCool Cyan-Blue\n#0087BD\n\n"
      "유진 : 전신·노트북·AI 아이콘\n모두 cyan-blue 컬러\n\n"
      "디지털·AI 세계를\n상징하는 차가운 빛\n\n"
      "Laptop screen glow :\ncyan-blue AI light on B&W room",
      size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

set_col_width(t_style, [5.5, 5.5, 5.5])
doc.add_paragraph()

# Gemini 공통 Selective Color 지시어
section_title(doc, "Gemini Selective Color Keyword Template", BLUE_DARK)
p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_after = Pt(6)
sf(p_kw.add_run(
    "Every Gemini prompt should include one of the following opening lines:\n\n"
    "For Sangwoo scenes :\n"
    '"Selective color photography/illustration: [scene description].\n'
    ' [Subject] rendered in warm golden amber color (#D48C00).\n'
    ' Everything else is desaturated black and white monochrome."\n\n'
    "For Yujin scenes :\n"
    '"Selective color photography/illustration: [scene description].\n'
    ' [Subject] rendered in cool cyan-blue color (#0087BD).\n'
    ' Everything else is desaturated black and white monochrome."\n\n'
    "For both characters together :\n"
    '"Selective color: Sangwoo rendered entirely in warm golden amber (#D48C00),\n'
    ' Yujin rendered entirely in cool cyan-blue (#0087BD).\n'
    ' Background and environment fully desaturated black and white monochrome."'
), size=10, color=GRAY_DARK)

doc.add_paragraph()

# CapCut 색 보정 팁
section_title(doc, "CapCut Selective Color Tips  /  편집 적용 방법", BLUE_DARK)
t_cap = doc.add_table(rows=4, cols=2)
t_cap.style = "Table Grid"
rows_cap = [
    ("Step 1", "Import all Runway video clips into CapCut timeline"),
    ("Step 2", "Apply 'Black & White' filter to each clip at 100% strength"),
    ("Step 3",
     "Add color adjustment overlay layer on top of each clip:\n"
     "- For Sangwoo clips : paint mask over Sangwoo's figure / set hue to amber range\n"
     "- For Yujin clips : paint mask over Yujin's figure / set hue to cyan-blue range"),
    ("Step 4",
     "Fine-tune saturation and brightness of colored areas:\n"
     "- Sangwoo amber : Saturation +40~60, Brightness +10\n"
     "- Yujin cyan : Saturation +40~60, Brightness +5\n"
     "- B&W areas : Saturation -100"),
]
for ri, (s, v) in enumerate(rows_cap):
    bg = "F0F4F8" if ri % 2 == 0 else "FFFFFF"
    shade(t_cap.rows[ri].cells[0], "1A5376")
    shade(t_cap.rows[ri].cells[1], bg)
    cfont(t_cap.rows[ri].cells[0], s, size=10, bold=True, color=WHITE)
    cfont(t_cap.rows[ri].cells[1], v, size=10)
set_col_width(t_cap, [2.5, 14.5])


# ══════════════════════════════════════════════════════════
#  Save
# ══════════════════════════════════════════════════════════
OUTPUT = (
    r"c:\Users\YSW1\Desktop\20260601 codyssey_mission"
    r"\N_B1-2_Contents\01_document\character_sheet.docx"
)
doc.save(OUTPUT)
print(f"Saved : {OUTPUT}")
